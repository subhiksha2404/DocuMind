import os
import uuid
import tempfile
import datetime
import json
import re
import zipfile
import time
import logging
import hashlib
import asyncio
import requests
import aiohttp
from collections import Counter
from config import Config
from typing import List,Dict,Optional
from fastapi import FastAPI, UploadFile, HTTPException, File, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from sentence_transformers import SentenceTransformer
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
import pandas as pd
from local_model_manager import model_manager

from chromadb import PersistentClient
from chromadb.utils import embedding_functions

from agents.orchestrator import AgentOrchestrator
from agents.langgraph_orchestrator import LangGraphOrchestrator
# Import LangChain text splitter
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    print("LangChain not available. Using fallback text splitting.")

# Initialize logging
logging.basicConfig(
    filename="app.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Initialize FastAPI
app = FastAPI()

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Your React frontend
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- GLOBAL PERSISTENT STORAGE PATHS --------------

DB_BASE_DIR = Config.VECTORSTORE_PATH

CHROMA_DB_DIR = Config.CHROMA_DB_DIR
LOG_FILE_PATH = Config.LOG_FILE_PATH
EMBEDDING_MODELS = Config.EMBEDDING_MODELS
INFERENCE_MODELS = Config.INFERENCE_MODELS

# Ensure required directories exist
os.makedirs(DB_BASE_DIR, exist_ok=True)
os.makedirs(CHROMA_DB_DIR, exist_ok=True)

agent_orchestrator = AgentOrchestrator()
langgraph_orchestrator = LangGraphOrchestrator()

# ---------- WebSocket Manager for Progress Updates ----------
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        self.active_connections.remove(websocket)

    async def send_progress(self, message: Dict, websocket: WebSocket):
        try:
            await websocket.send_json(message)
        except:
            self.disconnect(websocket)

manager = ConnectionManager()

# ---------- Initialize Persistent ChromaDB ---------------
print("🔧 PATH VERIFICATION:")
print(f"DB_BASE_DIR: {DB_BASE_DIR}")
print(f"CHROMA_DB_DIR: {CHROMA_DB_DIR}")
print(f"CHROMA_DB_DIR exists: {os.path.exists(CHROMA_DB_DIR)}")

client = PersistentClient(path=CHROMA_DB_DIR)
collection = client.get_or_create_collection("documents")


current_embedding_model = "sentence-transformers/all-MiniLM-L6-v2"
model = SentenceTransformer(current_embedding_model)

current_inference_model = "gemini"
# ---------- Load Embedding Model -------------------------
#model = SentenceTransformer("all-MiniLM-L6-v2")

# ------------------- HELPER FUNCTIONS ----------------------

async def update_progress(stage: str, progress: int, message: str = ""):
    """Broadcast progress update to all connected websockets"""
    for ws in manager.active_connections:
        try:
            await ws.send_json({
                "stage": stage,
                "progress": progress,
                "message": message,
                "timestamp": datetime.datetime.now().isoformat()
            })
        except:
            manager.disconnect(ws)


def get_file_hash(file_path):
    """Generate MD5 hash of file content for duplicate detection."""
    with open(file_path, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()

def safe_encode(texts, max_retries=3):
    """Safe embedding function with retries and error handling."""
    for attempt in range(max_retries):
        try:
            return model.encode(texts)
        except Exception as e:
            logging.error(f"Embedding failed (attempt {attempt + 1}): {e}")
            if attempt < max_retries - 1:
                time.sleep(1)  # Wait before retrying
            else:
                raise Exception(f"Embedding failed after {max_retries} retries: {e}")

def extract_pdf_metadata(file_path):
    """Extract metadata from PDF files."""
    try:
        reader = PdfReader(file_path)
        metadata = {
            "title": reader.metadata.get("/Title", "Untitled"),
            "author": reader.metadata.get("/Author", "Unknown"),
            "pages": len(reader.pages),
            "creation_date": str(reader.metadata.get("/CreationDate", "")),
            "modification_date": str(reader.metadata.get("/ModDate", ""))
        }
        return metadata
    except Exception as e:
        logging.error(f"PDF metadata extraction failed: {e}")
        return {"title": "Unknown", "author": "Unknown", "pages": 0}

def extract_docx_metadata(file_path):
    """Extract metadata from DOCX files."""
    try:
        doc = Document(file_path)
        # DOCX metadata is limited compared to PDF
        return {
            "title": "Document",  # DOCX doesn't have standard metadata fields
            "author": "Unknown",
            "pages": len(doc.paragraphs) // 50  # Rough estimate
        }
    except Exception as e:
        logging.error(f"DOCX metadata extraction failed: {e}")
        return {"title": "Unknown", "author": "Unknown", "pages": 0}

def extract_file_metadata(file_path, filename):
    """Extract metadata based on file type."""
    ext = filename.split(".")[-1].lower()
    
    if ext == "pdf":
        return extract_pdf_metadata(file_path)
    elif ext == "docx":
        return extract_docx_metadata(file_path)
    else:
        # For other file types, return basic metadata
        return {
            "title": filename,
            "author": "Unknown",
            "pages": 0
        }

# ------------------- PARSERS ----------------------
def parse_pdf(file_path):
    reader = PdfReader(file_path)
    return "\n\n".join([page.extract_text() or '' for page in reader.pages])

def parse_docx(file_path):
    doc = Document(file_path)
    return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def parse_txt(file_path):
    """Parse TXT files and extract text."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def parse_csv(file_path):
    try:
        df = pd.read_csv(file_path)
        return "\n".join(df.astype(str).apply(lambda x: " | ".join(x), axis=1))
    except Exception as e:
        logging.error(f"CSV parsing failed: {e}")
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()


def split_text_into_chunks(text):
    """Split text into semantic chunks."""
    if LANGCHAIN_AVAILABLE:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=512,  
            chunk_overlap=50,
            separators=["\n\n", "\n", " ", ""]
        )
        return splitter.split_text(text)
    else:
        return [p.strip() for p in text.split('\n\n') if p.strip()]

## ------------------- FILE PROCESSING WITH PROGRESS ----------------------

async def process_file_with_progress(file_path: str, filename: str):
    """Process file with progress updates."""

    print(f"📁 Processing {filename}")
    print(f"📁 Target ChromaDB: {CHROMA_DB_DIR}")
    print(f"📁 ChromaDB exists: {os.path.exists(CHROMA_DB_DIR)}")
    
    # Stage 1: File validation and metadata extraction
    await update_progress("validating", 10, "Validating file...")
    
    file_hash = get_file_hash(file_path)
    existing_files = get_collection().get(where={"file_hash": file_hash})
    
    if existing_files and existing_files["ids"]:
        raise Exception(f"File already indexed: {filename}")

    # Stage 2: Metadata extraction
    await update_progress("metadata", 20, "Extracting metadata...")
    extracted_metadata = extract_file_metadata(file_path, filename)
    
    # Stage 3: File parsing
    await update_progress("parsing", 30, "Parsing document content...")
    ext = filename.split(".")[-1].lower()
    
    if ext == "pdf":
        text = parse_pdf(file_path)
    elif ext == "docx":
        text = parse_docx(file_path)
    elif ext == "csv":
        text = parse_csv(file_path)
    elif ext == "txt":
        text = parse_txt(file_path)
    else:
        raise Exception(f"Unsupported file type: {filename}")

    if not text.strip():
        raise Exception(f"No readable text found in {filename}")

    # Stage 4: Text chunking
    await update_progress("chunking", 50, f"Splitting text into chunks...{filename}")
    chunks = split_text_into_chunks(text)
    
    if not chunks:
        raise Exception(f"No chunks created from {filename}")

    # Stage 5: Embedding generation
    await update_progress("embedding", 70, f"Generating embeddings...{filename}")
    embeddings = safe_encode(chunks).tolist()

    # Stage 6: Preparing metadata
    await update_progress("preparing", 85, f"Preparing {filename} for storage...")
    metadatas = []
    for chunk in chunks:
        chunk_metadata = {
            "filename": filename,
            "file_hash": file_hash,
            "chunk_length": len(chunk),
            **extracted_metadata
        }
        metadatas.append(chunk_metadata)

    # Stage 7: Storing in ChromaDB
    await update_progress("storing", 95, f"Storing {filename} in database...")

    collection = get_collection()
    print(f"📊 Collection before add: {collection.count()} vectors")

    ids = [str(uuid.uuid4()) for _ in chunks]
    get_collection().add(
        ids=ids,
        embeddings=embeddings,
        documents=chunks,
        metadatas=metadatas
    )

    print(f"📊 Collection after add: {collection.count()} vectors")
    print(f"✅ Stored {len(chunks)} chunks for {filename}")

    # Stage 8: Finalizing
    await update_progress("complete", 100, "Processing complete!")
    
    # Log the upload
    log_upload(filename, len(chunks), extracted_metadata, file_hash)
    
    return {
        "filename": filename,
        "chunks_stored": len(chunks),
        "embedding_dimension": len(embeddings[0]),
        "metadata": extracted_metadata
    }

def log_upload(filename, chunks_count, metadata, file_hash):
    log_data = []
    if os.path.exists(LOG_FILE_PATH):
        try:
            with open(LOG_FILE_PATH, "r") as f:
                log_data = json.load(f)
        except json.JSONDecodeError:
            log_data = []

    log_data.append({
        "filename": filename,
        "chunks_stored": chunks_count,
        "upload_time": str(datetime.datetime.now()),
        "metadata": metadata,
        "file_hash": file_hash
    })

    with open(LOG_FILE_PATH, "w") as f:
        json.dump(log_data, f, indent=4)

def get_collection():
    """Get or create collection with current embedding model."""
    return client.get_or_create_collection(
        name="documents",
        metadata={"embedding_model": current_embedding_model}
    )


async def generate_rag_response(query: str, context_chunks: List[str], conversation_history: List[Dict] = None):
    """Enhanced RAG response generation with adaptive prompting."""
    try:
        # Combine context chunks with better formatting
        context = "\n\n".join([f"📄 Source {i+1}:\n{chunk}" for i, chunk in enumerate(context_chunks)])
        
        # Adaptive prompt based on query type
        if any(word in query.lower() for word in ['what is', 'define', 'definition']):
            prompt = f"""Based EXCLUSIVELY on the provided document context, provide a comprehensive answer to the user's question.

DOCUMENT CONTEXT:
{context}

USER QUESTION: {query}

Please structure your answer as:
- **Clear Definition**: Start with a concise definition
- **Key Characteristics**: Main features or aspects
- **Importance/Significance**: Why it matters
- **Applications/Examples**: Where and how it's used
- **Summary**: Brief recap

Use markdown formatting with headers (##), bullet points, and **bold** for key terms.
Answer using ONLY the document context above. If the context doesn't contain enough information, say so."""
        
        elif any(word in query.lower() for word in ['how', 'process', 'steps', 'method']):
            prompt = f"""Based EXCLUSIVELY on the provided document context, explain the process or method asked about.

DOCUMENT CONTEXT:
{context}

USER QUESTION: {query}

Please structure your answer as:
- **Overview**: Brief introduction to the process
- **Step-by-Step Explanation**: Clear sequential steps
- **Key Components**: Important elements involved
- **Applications**: Where this process is used
- **Considerations**: Important factors or limitations

Use markdown formatting and be practical. Answer using ONLY the document context above."""
        
        elif any(word in query.lower() for word in ['compare', 'difference', 'vs', 'versus']):
            prompt = f"""Based EXCLUSIVELY on the provided document context, compare the concepts asked about.

DOCUMENT CONTEXT:
{context}

USER QUESTION: {query}

Please structure your answer as:
- **Overview**: Brief introduction to both concepts
- **Key Differences**: Clear comparison points
- **Similarities**: Common aspects
- **Use Cases**: When to use each
- **Summary**: Comparative conclusion

Use markdown formatting with tables or clear sections. Answer using ONLY the document context above."""
        
        else:
            # General purpose prompt for other queries
            prompt = f"""Based EXCLUSIVELY on the provided document context, provide a comprehensive and well-structured answer to the user's question.

DOCUMENT CONTEXT:
{context}

USER QUESTION: {query}

Please provide a natural, flowing answer that:
- Directly addresses the question
- Explains key concepts clearly
- Uses appropriate examples from the context
- Highlights important points
- Maintains logical flow

Use markdown formatting with:
## Headers for main sections
- Bullet points for lists
**Bold** for key terms and definitions
Clear paragraph structure

Answer using ONLY the document context above. If the information is insufficient, acknowledge this politely."""

        # Use local model manager
        answer = model_manager.generate_response(prompt)
        
        return answer
            
    except Exception as e:
        logging.error(f"RAG response generation failed: {e}")
        return f"I found relevant documents but encountered an error: {str(e)}"


# ------------------- ROUTES ----------------------

@app.websocket("/ws/progress")
async def websocket_progress(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            # Keep connection alive
            await websocket.receive_text()
    except WebSocketDisconnect:
        manager.disconnect(websocket)

@app.post("/set-embedding-model")
async def set_embedding_model(model_name: str):
    global current_embedding_model, model
    if model_name not in EMBEDDING_MODELS:
        raise HTTPException(status_code=400, detail="Invalid embedding model")
    
    current_embedding_model = model_name
    model = SentenceTransformer(model_name)
    
    # Recreate collection with new model
    global client
    client.delete_collection("documents")
    get_collection()
    
    return {"status": "success", "model": model_name}

@app.get("/embedding-models")
async def get_embedding_models():
    return {"models": list(EMBEDDING_MODELS.keys())}

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    """Upload and process a single file."""
    temp_dir = tempfile.mkdtemp()
    try:
        temp_path = os.path.join(temp_dir, file.filename)
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)

        result = await process_file_with_progress(temp_path, file.filename)
        
        return JSONResponse({"status": "success", **result})

    except Exception as e:
        logging.error(f"Upload failed: {e}")
        await update_progress("error", 0, f"Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)

@app.get("/status")
def status():
    count = get_collection().count()
    return {
        "total_vectors_stored": count,
        "database_path": CHROMA_DB_DIR,
        "embedding_model": current_embedding_model,
        "available_models": list(EMBEDDING_MODELS.keys())
    }


@app.delete("/delete-document")
async def delete_document(filename: str):
    try:
        results = collection.get(where={"filename": filename})
        if results and results["ids"]:
            collection.delete(ids=results["ids"])

        if os.path.exists(LOG_FILE_PATH):
            with open(LOG_FILE_PATH, "r") as f:
                log_data = json.load(f)
            log_data = [entry for entry in log_data if entry["filename"] != filename]
            with open(LOG_FILE_PATH, "w") as f:
                json.dump(log_data, f, indent=4)

        return {"status": "deleted", "filename": filename}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Add to your existing routes in main.py

@app.post("/search")
async def search(query: str, filter_author: str = None, filter_title: str = None, n_results: int = 10):
    """Search for documents with optional metadata filtering."""
    try:
        print(f"Search request - query: '{query}', author: '{filter_author}', title: '{filter_title}', n_results: {n_results}")
        
        # Create where clause for metadata filtering
        where_clause = {}
        if filter_author:
            where_clause["author"] = filter_author
        if filter_title:
            where_clause["title"] = filter_title
        
        # Generate query embedding
        query_embedding = safe_encode([query]).tolist()
        
        # Get collection count for debugging
        collection = get_collection()
        total_docs = collection.count()
        print(f"Total documents in collection: {total_docs}")
        
        # Query ChromaDB
        results = collection.query(
            query_embeddings=query_embedding,
            n_results=n_results,
            where=where_clause if where_clause else None
        )
        
        print(f"Search found {len(results['documents'][0]) if results['documents'] else 0} results")
        
        return {
            "query": query,
            "results": results
        }
        
    except Exception as e:
        logging.error(f"Search failed: {e}")
        print(f"Search error: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Search failed: {str(e)}"}
        )
    
@app.post("/upload-folder")
async def upload_folder(files: List[UploadFile] = File(...)):
    """Upload and process multiple files with progress updates."""
    results, errors = [], []
    temp_dir = tempfile.mkdtemp()

    try:
        total_files = len(files)

        if total_files == 0:
            # No files received from client
            await update_progress("error", 0, "No files received for folder upload.")
            return JSONResponse({
                "status": "failed",
                "documents_processed": [],
                "errors": [{"filename": "folder", "error": "No files received"}],
                "processed_count": 0,
                "error_count": 1,
                "total_files": 0
            })
        
        #starting
        await update_progress("saving", 0, f"Starting upload of {total_files} files...")

        print(f"Received {total_files} files")

        for i, file in enumerate(files):
            try:
                temp_path = os.path.join(temp_dir, file.filename)
                print(f"Saving file: {file.filename}")  

                with open(temp_path, "wb") as f:
                    content = await file.read()
                    f.write(content)

                # Broadcast saving progress
                await update_progress("saving",int((i + 1) / total_files * 20),f"Saving files... ({i+1}/{total_files})")

                # Progress: processing start
                await update_progress("processing",int((i + 1) / total_files * 40),f"Processing {file.filename} ({i+1}/{total_files})"
                )

                # Process file with progress updates
                result = await process_file_with_progress(temp_path, file.filename)
                results.append(result)

                # Progress: processed
                await update_progress(
                    "processing",
                    int((i + 1) / total_files * 85),
                    f"Finished processing {file.filename} ({i+1}/{total_files})"
                )

            except Exception as e:
                error_msg = str(e)
                errors.append({"filename": file.filename, "error": error_msg})
                logging.error(f"Error processing {file.filename}: {e}")
                await update_progress("error",int((i + 1) / total_files * 100), f"{file.filename}: {error_msg}")

        # Final broadcast
        await update_progress("complete",100,f"Folder processing complete! {len(results)} successful, {len(errors)} errors")

        return JSONResponse({
            "status": "success" if results else "failed",
            "documents_processed": results,
            "errors": errors,
            "processed_count": len(results),
            "error_count": len(errors),
            "total_files": total_files
        })

    except Exception as e:
        logging.error(f"Folder upload failed(unexpected): {e}")
        return JSONResponse({
            "status": "failed",
            "documents_processed": results,
            "errors": [{"filename": "folder", "error": str(e)}],
            "processed_count": len(results),
            "error_count": len(errors) + 1,
            "total_files": len(files) if files else 0
        })

    finally:
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)


@app.get("/uploaded-files")
async def get_uploaded_files():
    """Get list of all uploaded files with metadata."""
    try:
        if not os.path.exists(LOG_FILE_PATH):
            return {"files": []}
        
        with open(LOG_FILE_PATH, "r") as f:
            log_data = json.load(f)
        
        return {"files": log_data}
    except Exception as e:
        logging.error(f"Failed to get uploaded files: {e}")
        return {"files": []}
    

@app.post("/set-inference-model")
async def set_inference_model(model_name: str):
    """Change the current inference model."""
    global current_inference_model
    
    # Update available models configuration
    INFERENCE_MODELS.update({
        "gemini": {"api_url": "local", "requires_key": False},
        "microsoft/DialoGPT-large": {"api_url": "local", "requires_key": False},
        "microsoft/DialoGPT-medium": {"api_url": "local", "requires_key": False}, 
        "google/flan-t5-base": {"api_url": "local", "requires_key": False}
    })
    
    if model_name not in INFERENCE_MODELS:
        raise HTTPException(status_code=400, detail="Invalid inference model")
    
    # Load the model
    model_manager.load_local_model(model_name)
    current_inference_model = model_name
    
    return {"status": "success", "model": model_name}

# Update the get_inference_models route:
@app.get("/inference-models")
async def get_inference_models():
    """Get available inference models."""
    models = {
        "gemini": {"api_url": "local", "requires_key": bool(Config.GEMINI_API_KEY)},
        "microsoft/DialoGPT-large": {"api_url": "local", "requires_key": False},
        "microsoft/DialoGPT-medium": {"api_url": "local", "requires_key": False},
        "google/flan-t5-base": {"api_url": "local", "requires_key": False}
    }
    return {"models": list(models.keys())}

@app.post("/chat")
async def chat_with_documents(query: str, conversation_id: Optional[str] = None, n_context_chunks: int = 5):
    """Chat with documents using RAG."""
    try:
        # Use the agent orchestrator (same functionality, better organization)
        response_data = agent_orchestrator.process_query(query, n_context_chunks)
        
        # Add the current model info (same as before)
        response_data["model_used"] = current_inference_model
        
        return response_data
        
    except Exception as e:
        logging.error(f"Chat failed: {e}")
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")
    
@app.post("/simple-chat")
async def simple_chat_with_documents(query: str, conversation_id: Optional[str] = None, n_context_chunks: int = 5):
    """Original simple RAG implementation as fallback."""
    try:
        # This is your EXACT original chat_with_documents code
        # Step 1: Semantic search to find relevant chunks
        query_embedding = safe_encode([query]).tolist()
        collection = get_collection()
        
        search_results = collection.query(
            query_embeddings=query_embedding,
            n_results=n_context_chunks
        )
        
        # Extract context chunks
        context_chunks = []
        if search_results and search_results['documents']:
            context_chunks = search_results['documents'][0]
        
        # Step 2: Generate RAG response
        answer = await generate_rag_response(query, context_chunks)
        
        # Step 3: Return response with sources
        response_data = {
            "question": query,
            "answer": answer,
            "sources": [],
            "model_used": current_inference_model
        }
        
        # Add source information
        if search_results and search_results['metadatas']:
            seen_sources = set()
            for metadata in search_results['metadatas'][0]:
                source_key = f"{metadata.get('filename', '')}-{metadata.get('title', '')}"
                if source_key not in seen_sources:
                    seen_sources.add(source_key)
                    source_info = {
                        "filename": metadata.get('filename', 'Unknown'),
                        "title": metadata.get('title', 'Untitled'),
                        "author": metadata.get('author', 'Unknown'),
                    }
                    response_data["sources"].append(source_info)
        
        return response_data
        
    except Exception as e:
        logging.error(f"Simple chat failed: {e}")
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")
    
@app.post("/chat-langgraph")
async def chat_with_documents_langgraph(query: str, conversation_id: Optional[str] = None, n_context_chunks: int = 5):
    """Chat with documents using LangGraph agentic RAG."""
    try:
        response_data = langgraph_orchestrator.process_query(query, n_context_chunks)
        response_data["model_used"] = current_inference_model
        return response_data
    except Exception as e:
        logging.error(f"LangGraph chat failed: {e}")
        # Fallback to custom agents
        return await chat_with_documents(query, conversation_id, n_context_chunks)
    
@app.get("/debug-chroma")
async def debug_chroma():
    """Debug ChromaDB state"""
    try:
        collection = get_collection()
        count = collection.count()
        
        # Get some sample data
        sample = collection.get(limit=5)
        
        return {
            "chroma_db_path": CHROMA_DB_DIR,
            "collection_count": count,
            "sample_documents": sample.get('documents', []),
            "sample_metadatas": sample.get('metadatas', []),
            "available_collections": [col.name for col in client.list_collections()]
        }
    except Exception as e:
        return {"error": str(e)}